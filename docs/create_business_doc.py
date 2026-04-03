#!/usr/bin/env python3
"""
Generate DiamondBets business overview as Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DiamondBets Business Overview', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
doc.add_heading('Executive Summary', 1)
p = doc.add_paragraph()
p.add_run('DiamondBets').bold = True
p.add_run(' is a systematic sports betting platform that combines machine learning predictions with automated bankroll management. Users receive high-confidence MLB picks with proven 65% accuracy, automated bet sizing via Kelly Criterion, and seamless integration with ESPN Bet for one-click execution.')

doc.add_page_break()

# Product Overview
doc.add_heading('Product Overview', 1)

doc.add_heading('What We Do', 2)
items = [
    'Generate daily MLB predictions using ensemble ML model (60% expert system + 40% gradient boosting)',
    'Filter to premium picks only (64%+ win probability with measurable edge vs Vegas)',
    'Calculate optimal bet sizing based on user\'s bankroll and risk tolerance',
    'Track performance automatically (results, P/L, ROI)',
    'Enable one-click betting via ESPN Bet API integration'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Target Customer', 2)
doc.add_paragraph('Quantitatively-minded sports bettors who want systematic edge detection without manual calculation or guesswork.')

doc.add_page_break()

# How It Works
doc.add_heading('How It Works', 1)

doc.add_heading('1. Prediction Engine (Backend - Runs Daily at 10 AM ET)', 2)

doc.add_heading('Data Sources:', 3)
sources = [
    'MLB Stats API (team/pitcher stats, schedules)',
    'ESPN odds API (market lines, implied probability)',
    'Historical game results (2025 season backtest: 2,426 games)'
]
for source in sources:
    doc.add_paragraph(source, style='List Bullet')

doc.add_heading('Model Architecture:', 3)
doc.add_paragraph('Ensemble Prediction = (Expert System × 60%) + (ML Model × 40%)')

p = doc.add_paragraph()
p.add_run('Expert System:\n').bold = True
expert_items = [
    'Pitcher quality (24% weight): ERA, WHIP, K/9, recent form',
    'Team offense (16%): OPS, runs/game',
    'Bullpen strength (10%): ERA, saves',
    'Home field advantage',
    'Pythagorean win percentage'
]
for item in expert_items:
    doc.add_paragraph(item, style='List Bullet 2')

p = doc.add_paragraph()
p.add_run('ML Model:\n').bold = True
ml_items = [
    'Logistic regression + gradient boosting',
    'Trained on 2,426 games',
    'Automatically learns optimal feature weights'
]
for item in ml_items:
    doc.add_paragraph(item, style='List Bullet 2')

doc.add_heading('Quality Filter:', 3)
filter_items = [
    'Only publish picks ≥64% win probability',
    'Must show positive edge vs Vegas implied probability',
    'Result: 2-4 premium picks per day (some days zero)'
]
for item in filter_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('2. User Dashboard (Frontend - Real-Time)', 2)

doc.add_heading('First-Time Setup:', 3)
setup = [
    'User sets starting bankroll (e.g., $1,000)',
    'Chooses Kelly fraction (Conservative: 1/4, Moderate: 1/2, Aggressive: Full)',
    'Optionally connects ESPN Bet account for automation'
]
for item in setup:
    doc.add_paragraph(item, style='List Number')

doc.add_heading('Daily Workflow:', 3)
workflow = [
    '10 AM ET → Picks published',
    'Dashboard shows each pick with: Team + matchup, Model probability (e.g., 68% Dodgers), Edge vs market (e.g., +4.2%), Recommended bet size (e.g., $47.50) ← Auto-calculated via Kelly',
    'User clicks "Place Bet"',
    'Manual Mode: Logs bet, deducts from bankroll, marks pending',
    'Automation Mode: API places bet via ESPN Bet, auto-logs',
    'End of day → Results scraped from MLB API',
    'Bets auto-marked Win/Loss',
    'Profit calculated, bankroll updated',
    'P/L and stats refreshed'
]
for item in workflow:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Key Features:', 3)
features = [
    'Auto bet sizing: Kelly Criterion runs automatically, no manual calculation',
    'Bankroll tracking: Current balance, total P/L, win/loss record',
    'Performance analytics: Win rate, ROI, bet history',
    'Automation (ESPN Bet): One-click execution, auto result tracking'
]
for item in features:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

doc.add_heading('3. Monetization', 2)

p = doc.add_paragraph()
p.add_run('Revenue Model: Subscription + Affiliate\n\n').bold = True

doc.add_heading('Free Tier', 3)
free = [
    'Public site only (read-only picks)',
    'See today\'s premium picks + track record',
    'No dashboard access',
    'No bet tracking or analytics'
]
for item in free:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Paid Tier ($40/month)', 3)
paid = [
    'Full dashboard access',
    'Bankroll tracking + Kelly calculator',
    'Auto bet sizing on every pick',
    'Manual bet logging',
    'Performance analytics (P/L, win rate, ROI)',
    'Bet history'
]
for item in paid:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Automation Feature (Free with ESPN Bet)', 3)
automation = [
    'Same $40/month subscription',
    'Connect ESPN Bet account → automation unlocks at no extra cost',
    'One-click betting via API',
    'Auto result tracking (no manual W/L marking)',
    'Live odds sync',
    'Drives affiliate conversion without additional charge'
]
for item in automation:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run('\nThe Pitch: ').bold = True
p.add_run('"Want automation? Just connect your ESPN Bet account — no extra charge."')

doc.add_heading('Unit Economics:', 3)
economics = [
    'User subscribes → $40/month (dashboard access)',
    'User connects ESPN Bet for automation → $150 CPA (affiliate)',
    'First-month revenue: $40 subscription + $150 affiliate = $190',
    'Ongoing: $40/month recurring'
]
for item in economics:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Assumptions:', 3)
assumptions = [
    '40% of paid users connect ESPN Bet for automation',
    '60% retention after month 1',
    'Average LTV (12 months): ($40 × 7 months avg) + ($150 × 40% automation rate) = $280 + $60 = $340'
]
for item in assumptions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Technical Architecture
doc.add_heading('Technical Architecture', 1)

doc.add_heading('Stack', 2)
stack = [
    'Frontend: Vanilla HTML/CSS/JS (hosted on GitHub Pages)',
    'Backend: Node.js/Express (hosted on Railway)',
    'Database: PostgreSQL (Railway managed)',
    'Prediction engine: Python (runs daily via cron on Mac Mini)',
    'APIs: MLB Stats API, ESPN odds, ESPN Bet (for automation)'
]
for item in stack:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Data Flow', 2)
flow = [
    'Daily 10 AM: Python predictor → Generates picks.json → Uploaded to GitHub',
    'Dashboard fetches picks.json → Calculates Kelly bets → Displays to user',
    'User places bet → Logged to PostgreSQL (backend API)',
    'End of day: Python scraper → Fetches game results → Updates results.json',
    'Dashboard checks results → Marks bets W/L → Updates P/L'
]
for item in flow:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Competitive Advantages
doc.add_heading('Competitive Advantages', 1)

doc.add_heading('1. Proven Performance', 2)
perf = [
    '65% win rate on premium picks (vs 52% industry standard)',
    '+24% ROI (backtested on 2,426 games)',
    'Full transparency: every pick, every result, public track record'
]
for item in perf:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2. Automation > Manual Picks', 2)
auto = [
    'Not selling picks — selling a system',
    'No daily grind of posting picks',
    'Tool-first positioning (vs guru/tipster)'
]
for item in auto:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('3. Automation Drives Affiliate Conversion', 2)
affiliate = [
    'Automation unlocks with ESPN Bet (no extra charge)',
    'Feels like a feature unlock, not a paywall',
    'Product becomes the affiliate funnel',
    '40%+ of paid users connect ESPN Bet for automation'
]
for item in affiliate:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4. Recurring Revenue', 2)
recurring = [
    'Subscription model > one-time sales',
    'Predictable MRR growth',
    'Lower churn (users invested in bankroll tracking)'
]
for item in recurring:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Go-To-Market
doc.add_heading('Go-To-Market Plan', 1)

doc.add_heading('Phase 1: MVP Launch (Week 1-2)', 2)
phase1 = [
    '✅ Public site (free picks, track record)',
    '✅ Dashboard v1 (manual mode with Kelly calc)',
    '⏳ ESPN Bet affiliate signup',
    '⏳ Backend deployment (Railway)'
]
for item in phase1:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Phase 2: Automation (Week 3-4)', 2)
phase2 = [
    'ESPN Bet API integration',
    'One-click betting',
    'Auto result tracking',
    'Paid tier launch ($40/month)'
]
for item in phase2:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Phase 3: Growth (Month 2+)', 2)
phase3 = [
    'SEO content (betting guides, model explanation)',
    'Reddit/Twitter organic marketing',
    'Discord community for users',
    'Referral program (share revenue with power users)'
]
for item in phase3:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Financial Projections
doc.add_heading('Financial Projections (12-Month)', 1)

doc.add_heading('Conservative Scenario:', 2)
conservative = [
    '200 free users/month (organic traffic to public site)',
    '5% convert to paid dashboard → 10 paid users/month',
    '40% of paid users connect ESPN Bet → 4 affiliate conversions/month',
    'Month 1 revenue: ($40 × 10 subscribers) + ($150 × 4 CPA) = $400 + $600 = $1,000',
    'Month 12: ($40 × 70 subscribers) + ($150 × 4 new CPA) = $2,800 + $600 = $3,400/month'
]
for item in conservative:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Moderate Scenario (with marketing):', 2)
moderate = [
    '1,000 free users/month (SEO + Reddit/Twitter marketing)',
    '8% convert to paid → 80 paid users/month',
    '40% connect ESPN Bet → 32 affiliate conversions/month',
    'Month 1 revenue: ($40 × 80) + ($150 × 32) = $3,200 + $4,800 = $8,000',
    'Month 12: ($40 × 560 subscribers) + ($150 × 32 CPA) = $22,400 + $4,800 = $27,200/month'
]
for item in moderate:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Risks
doc.add_heading('Risks & Mitigation', 1)

table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

# Header row
header_cells = table.rows[0].cells
header_cells[0].text = 'Risk'
header_cells[1].text = 'Mitigation'

# Data rows
risks = [
    ('Model accuracy drops', 'Backtest quarterly, retrain on new data, publish honest results'),
    ('ESPN Bet blocks API access', 'Pivot to manual mode, partner with different book'),
    ('Low conversion to paid', 'Improve free tier value, test pricing ($30 vs $50)'),
    ('Regulatory (gambling laws)', 'No financial advice, entertainment-only disclaimers, geo-fence if needed')
]

for i, (risk, mitigation) in enumerate(risks, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = risk
    row_cells[1].text = mitigation

doc.add_page_break()

# Next Steps
doc.add_heading('Next Steps', 1)

doc.add_heading('Immediate (This Week):', 2)
immediate = [
    'Finalize ESPN Bet affiliate approval',
    'Deploy backend to Railway',
    'Test automation flow end-to-end'
]
for i, item in enumerate(immediate, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Short-Term (Next 2 Weeks):', 2)
short_term = [
    'Launch paid tier',
    'Onboard first 10 paying users',
    'Collect feedback, iterate'
]
for i, item in enumerate(short_term, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Long-Term (3-6 Months):', 2)
long_term = [
    'Expand to NFL/NBA (new sports = new TAM)',
    'White-label licensing to betting communities',
    'API access for quant traders'
]
for i, item in enumerate(long_term, 1):
    doc.add_paragraph(f'{i}. {item}')

# Footer
doc.add_paragraph('\n')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Last updated: April 2, 2026')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save
doc.save('DiamondBets_Business_Overview.docx')
print('✅ Document created: DiamondBets_Business_Overview.docx')
