#!/usr/bin/env python3
"""
Generate DiamondBets business overview (Kalshi-focused) as Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DiamondBets Business Overview', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
doc.add_heading('Executive Summary', 1)
doc.add_paragraph('Algorithmic MLB trading platform for prediction markets. Combines proven ML model (65% accuracy, +24% ROI) with fully automated execution via Kalshi API. Target: quantitative traders seeking systematic edge in sports prediction markets.')

# Product
doc.add_heading('Product', 1)
doc.add_paragraph('Daily MLB predictions using ensemble model (60% expert + 40% ML). Auto-executes trades on Kalshi prediction markets via API. Users set bankroll + risk tolerance, system handles everything else. Full automation, zero fees, CFTC-regulated.')

# How It Works
doc.add_heading('How It Works', 1)

doc.add_heading('Prediction Engine', 2)
doc.add_paragraph('Runs daily at 10 AM ET. Analyzes pitcher stats, team metrics, park factors, market odds. Blends expert system (hand-tuned) with ML model (trained on 2,426 games). Filters to 64%+ probability with measurable edge vs Kalshi market price.')

doc.add_heading('User Flow', 2)
flow = [
    'User connects Kalshi account via API',
    'Sets bankroll ($500-10,000) and Kelly fraction (conservative to aggressive)',
    '10 AM: Model generates picks, calculates position sizes',
    'System auto-executes trades via Kalshi API',
    'End of day: Positions settled, P/L calculated, bankroll updated',
    'Dashboard shows performance: ROI, win rate, trade history'
]
for i, item in enumerate(flow, 1):
    doc.add_paragraph(f'{i}. {item}')

# Why Kalshi
doc.add_heading('Why Kalshi vs Traditional Sportsbooks', 1)
advantages = [
    'Full API: Real automated execution (not fake "automation")',
    '0% fees: No vig (vs 4-10% at sportsbooks)',
    'Explicitly allows bots: Many quant traders already using it',
    'CFTC regulated: Legitimate, won\'t shut down API access',
    'Better user base: Quant-minded traders (less likely to have accounts already)'
]
for item in advantages:
    doc.add_paragraph(item, style='List Bullet')

# Monetization
doc.add_heading('Monetization', 1)

doc.add_heading('Pricing Tiers', 2)

doc.add_paragraph('Free Tier: Public site only (read-only picks + track record)')

doc.add_paragraph('Standard ($50/month): Dashboard access, manual mode (Kelly calc, bet tracker, analytics). For users who want to execute trades manually.')

doc.add_paragraph('Pro ($75/month): Full automation via Kalshi API. One-click setup, hands-free execution, auto P/L tracking. Requires Kalshi account (drives referrals).')

doc.add_paragraph('API Access ($150/month): Direct API access to picks data. For quant traders building custom strategies. JSON endpoint with probabilities, edge calculations, historical data.')

doc.add_heading('Revenue Model', 2)
revenue = [
    'Subscription revenue: $50-150/month depending on tier',
    'Kalshi referrals: $25 per signup (lower than sportsbooks but acceptable)',
    'Target: 60% of Pro users are new to Kalshi (drives referral revenue)',
    'Average user value: $75/month subscription + ($25 × 60% referral rate) = $90 first month, then $75/month'
]
for item in revenue:
    doc.add_paragraph(item, style='List Bullet')

# Financial Projections
doc.add_heading('Financial Projections (12-Month)', 1)

table = doc.add_table(rows=3, cols=4)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Scenario'
header_cells[1].text = 'Paid Users/Month'
header_cells[2].text = 'Avg Tier Mix'
header_cells[3].text = 'Month 12 Revenue'

row1 = table.rows[1].cells
row1[0].text = 'Conservative'
row1[1].text = '50 new/month'
row1[2].text = '60% Pro, 30% Standard, 10% API'
row1[3].text = '$23,625/month'

row2 = table.rows[2].cells
row2[0].text = 'Moderate'
row2[1].text = '150 new/month'
row2[2].text = '60% Pro, 30% Standard, 10% API'
row2[3].text = '$70,875/month'

doc.add_paragraph()
doc.add_paragraph('Conservative calc (Month 12): (350 Pro × $75) + (175 Standard × $50) + (58 API × $150) = $26,250 + $8,750 + $8,700 = $43,700/month MRR')
doc.add_paragraph('Moderate calc (Month 12): (1,050 Pro × $75) + (525 Standard × $50) + (175 API × $150) = $78,750 + $26,250 + $26,250 = $131,250/month MRR')

# Competitive Advantages
doc.add_heading('Competitive Advantages', 1)
advantages = [
    'Only MLB algo trading platform on Kalshi (zero direct competition)',
    'Proven 65% win rate, +24% ROI with full transparency',
    'True automation (not browser scripts or manual deep links)',
    'Zero-fee execution (massive advantage vs sportsbook vig)',
    'Positioned as "trading" not "betting" (attracts quant audience)'
]
for item in advantages:
    doc.add_paragraph(item, style='List Bullet')

# Target Market
doc.add_heading('Target Market', 1)

doc.add_paragraph('Primary: Quantitative retail traders who already use platforms like QuantConnect, Alpaca, Interactive Brokers. Comfortable with APIs, automation, backtesting. Looking for diversification beyond stocks/crypto.')

doc.add_paragraph('Secondary: Sports analytics enthusiasts who understand edge/EV but haven\'t found a systematic way to trade it. More technical than typical sports bettors.')

doc.add_paragraph('NOT targeting: Casual sports bettors, "lock of the day" followers, parlay degenerates.')

# Marketing Strategy
doc.add_heading('Go-To-Market', 1)

doc.add_paragraph('Phase 1 (Week 1-2): Launch Standard tier ($50/month, manual mode). Prove product works. Collect testimonials.')

doc.add_paragraph('Phase 2 (Week 3-4): Build Kalshi API integration. Launch Pro tier ($75/month, automation). Onboard first 20 automated users.')

doc.add_paragraph('Phase 3 (Month 2+): API tier for quant traders ($150/month). Marketing: r/algotrading, QuantConnect forums, Kalshi Discord, finance Twitter.')

channels = [
    'Reddit: r/algotrading, r/sportsbetting, r/Kalshi',
    'Twitter/X: Algo trading community, sports analytics folks',
    'Discord: Kalshi official server, quant trading servers',
    'SEO: "Kalshi MLB trading bot", "automated sports prediction markets"'
]
for item in channels:
    doc.add_paragraph(item, style='List Bullet')

# Technical Stack
doc.add_heading('Technical Stack', 1)
stack = [
    'Frontend: React dashboard (not vanilla HTML for this audience)',
    'Backend: Python FastAPI (handles Kalshi API integration)',
    'Database: PostgreSQL (user accounts, trade logs, performance)',
    'Hosting: Railway (backend) + Vercel (frontend)',
    'Prediction engine: Python (existing model, runs daily)',
    'Kalshi integration: Official Python SDK'
]
for item in stack:
    doc.add_paragraph(item, style='List Bullet')

# Risks
doc.add_heading('Key Risks & Mitigation', 1)

table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Light Grid Accent 1'

header_cells2 = table2.rows[0].cells
header_cells2[0].text = 'Risk'
header_cells2[1].text = 'Mitigation'

risks = [
    ('Model accuracy drops', 'Quarterly backtests, retrain on new data, publish honest results. Subscription model means users stay if long-term edge holds.'),
    ('Kalshi changes API or fees', 'Diversify to Polymarket, PredictIt if needed. Standard/API tiers still work manually.'),
    ('Low market adoption', 'Start with proven model + free tier to build trust. API tier targets smaller but higher-value market.')
]

for i, (risk, mitigation) in enumerate(risks, 1):
    row_cells = table2.rows[i].cells
    row_cells[0].text = risk
    row_cells[1].text = mitigation

# Next Steps
doc.add_heading('Immediate Next Steps', 1)
next_steps = [
    'Build Standard tier dashboard (manual mode, Kelly calc, P/L tracking)',
    'Test Kalshi API integration (place test trades, verify execution)',
    'Launch landing page targeting "MLB algorithmic trading" keywords',
    'Price test: $50 vs $75 for Standard tier',
    'Onboard first 10 users (collect feedback, iterate)'
]
for i, item in enumerate(next_steps, 1):
    doc.add_paragraph(f'{i}. {item}')

# Footer
doc.add_paragraph('\n')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Last updated: April 2, 2026')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save
doc.save('DiamondBets_Kalshi_Business_Plan.docx')
print('✅ Kalshi business plan created: DiamondBets_Kalshi_Business_Plan.docx')
