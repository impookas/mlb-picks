#!/usr/bin/env python3
"""
Generate compact DiamondBets business overview as Word document.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('DiamondBets Business Overview', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Executive Summary
doc.add_heading('Executive Summary', 1)
doc.add_paragraph('Systematic MLB betting platform combining ML predictions (65% accuracy) with automated bankroll management. Users get high-confidence picks, Kelly-based bet sizing, and one-click execution via ESPN Bet integration.')

# Product
doc.add_heading('Product', 1)
doc.add_paragraph('Daily MLB predictions using ensemble model (60% expert system + 40% ML). Only publishes picks ≥64% probability with measurable edge vs Vegas. Auto-calculates optimal bet size, tracks performance, enables one-click betting.')

# How It Works
doc.add_heading('How It Works', 1)

doc.add_heading('Prediction Engine', 2)
doc.add_paragraph('Runs daily at 10 AM ET. Analyzes pitcher stats, team offense, bullpen, home field, park factors. Blends expert system (hand-tuned weights) with ML model (trained on 2,426 games). Filters to 64%+ confidence only.')

doc.add_heading('User Flow', 2)
flow = [
    '10 AM: Picks published',
    'Dashboard shows recommended bet size (Kelly calc)',
    'User clicks "Place Bet" (manual or auto via ESPN Bet)',
    'End of day: Results auto-tracked, P/L calculated'
]
for item in flow:
    doc.add_paragraph(item, style='List Number')

# Monetization
doc.add_heading('Monetization', 1)

doc.add_heading('Pricing Tiers', 2)
doc.add_paragraph('Free Tier: Public site (read-only picks + track record)')
doc.add_paragraph('Paid Tier ($40/month): Dashboard, bankroll tracking, Kelly calc, bet logging, analytics')
doc.add_paragraph('Automation (Free with ESPN Bet): One-click betting, auto result tracking. Unlocks when user connects ESPN Bet account (drives affiliate conversion).')

doc.add_heading('Revenue Model', 2)
revenue = [
    'Subscription: $40/month per paid user',
    'Affiliate: $150 CPA when user connects ESPN Bet for automation',
    'First-month value: $40 + ($150 × 40% automation rate) = $100 avg',
    'LTV (12 months): $340'
]
for item in revenue:
    doc.add_paragraph(item, style='List Bullet')

# Financial Projections
doc.add_heading('Financial Projections (12-Month)', 1)

table = doc.add_table(rows=3, cols=4)
table.style = 'Light Grid Accent 1'

header_cells = table.rows[0].cells
header_cells[0].text = 'Scenario'
header_cells[1].text = 'Free Users/Month'
header_cells[2].text = 'Paid Conversions'
header_cells[3].text = 'Month 12 Revenue'

row1 = table.rows[1].cells
row1[0].text = 'Conservative'
row1[1].text = '200'
row1[2].text = '5% (10/mo)'
row1[3].text = '$3,400/month'

row2 = table.rows[2].cells
row2[0].text = 'Moderate'
row2[1].text = '1,000'
row2[2].text = '8% (80/mo)'
row2[3].text = '$27,200/month'

# Competitive Advantages
doc.add_heading('Competitive Advantages', 1)
advantages = [
    'Proven 65% win rate (+24% ROI) with full transparency',
    'Tool-first positioning (not selling picks, selling a system)',
    'Automation drives ESPN Bet affiliate conversion (40% attach rate)',
    'Recurring revenue model (lower churn, predictable MRR)'
]
for item in advantages:
    doc.add_paragraph(item, style='List Bullet')

# Roadmap
doc.add_heading('Roadmap', 1)

doc.add_paragraph('Phase 1 (Week 1-2): ESPN Bet affiliate approval, backend deployment (Railway)')
doc.add_paragraph('Phase 2 (Week 3-4): ESPN Bet API integration, one-click betting, paid tier launch')
doc.add_paragraph('Phase 3 (Month 2+): SEO content, Reddit/Twitter marketing, referral program')

# Risks
doc.add_heading('Key Risks', 1)
risks = [
    'Model accuracy: Backtest quarterly, retrain on new data',
    'ESPN Bet API access: Pivot to manual mode if blocked',
    'Low conversion: Test pricing ($30-50 range), improve free tier value'
]
for item in risks:
    doc.add_paragraph(item, style='List Bullet')

# Next Steps
doc.add_heading('Immediate Next Steps', 1)
next_steps = [
    'Finalize ESPN Bet affiliate signup',
    'Deploy backend to Railway',
    'Test automation end-to-end',
    'Launch paid tier ($40/month)',
    'Onboard first 10 users'
]
for i, item in enumerate(next_steps, 1):
    doc.add_paragraph(f'{i}. {item}')

# Footer
doc.add_paragraph('\n')
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Last updated: April 2, 2026')
run.italic = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

# Save
doc.save('DiamondBets_Business_Overview_Compact.docx')
print('✅ Compact document created: DiamondBets_Business_Overview_Compact.docx')
